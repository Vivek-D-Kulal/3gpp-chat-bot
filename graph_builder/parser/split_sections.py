from docx import Document
import re

def is_section_heading(text):
    return re.match(r"^\d+(\.\d+)*(\s+.+)?$", text.strip())

def is_bullet(paragraph):
    style = paragraph.style.name.lower()
    text = paragraph.text.strip()
    return style.startswith("list") or text.startswith(("-", "•", "*", "▪"))

def normalize_section_id(raw):
    """Extracts only the numeric portion from a section heading (e.g., '8.2.20.5 HashMME' -> '8.2.20.5')"""
    return raw.strip().split()[0]

def split_docx_into_sections(doc_path):
    """Handles .docx files with table and bullet formatting."""
    doc = Document(doc_path)
    sections = {}
    current_section_id = None

    content_blocks = list(doc.paragraphs) + list(doc.tables)
    content_blocks.sort(key=lambda b: b._element.getparent().index(b._element))

    for block in content_blocks:
        if hasattr(block, "text"):  # Paragraph
            text = block.text.strip()
            if not text:
                continue

            if is_section_heading(text):
                parts = text.split(" ", 1)
                raw_id = parts[0]
                section_id = normalize_section_id(raw_id)
                section_title = parts[1] if len(parts) > 1 else ""
                current_section_id = section_id

                if current_section_id not in sections:
                    sections[current_section_id] = {
                        "title": section_title,
                        "content": [],
                        "tables": []
                    }
            else:
                if current_section_id:
                    entry_type = "bullet" if is_bullet(block) else "text"
                    sections[current_section_id]["content"].append({
                        "type": entry_type,
                        "text": text,
                        "style": block.style.name
                    })
        else:  # Table
            if current_section_id:
                table_data = []
                for row in block.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                sections[current_section_id]["tables"].append({"rows": table_data})

    return sections

def split_text_into_sections(text):
    """Fallback for plain-text `.doc` files (no bullets or tables)."""
    sections = {}
    current_section = None
    current_title = ""
    current_lines = []

    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue

        if is_section_heading(line):
            if current_section:
                sections[current_section] = {
                    "title": current_title,
                    "content": [{"type": "text", "text": l} for l in current_lines],
                    "tables": []
                }

            parts = line.split(" ", 1)
            raw_id = parts[0]
            current_section = normalize_section_id(raw_id)
            current_title = parts[1] if len(parts) > 1 else ""
            current_lines = []
        else:
            current_lines.append(line)

    # Add last section
    if current_section and current_lines:
        sections[current_section] = {
            "title": current_title,
            "content": [{"type": "text", "text": l} for l in current_lines],
            "tables": []
        }

    return sections
